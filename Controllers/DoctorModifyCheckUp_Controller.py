from PyQt5.QtWidgets import QMainWindow, QVBoxLayout, QCheckBox, QMessageBox
from Views.Doctor_Diagnosis import Ui_Doctor_Diagnosis as DoctorDiagnosisUI
from Models.CheckUp import CheckUp
from Models.Patient import Patient
from Models.Doctor import Doctor
from Models.LaboratoryTest import Laboratory
from datetime import datetime, date
from docx import Document
import os
from docx2pdf import convert

class DoctorDiagnosisModify(QMainWindow):
    def __init__(self, checkup_id, doc_id, parent=None):
        super().__init__(parent)
        self.ui = DoctorDiagnosisUI()
        self.ui.setupUi(self)

        # Store the checkup ID, doc_id, and reference to the DoctorDashboard
        self.checkup_id = checkup_id
        self.doc_id = doc_id

        print(f"DoctorDiagnosis initialized successfully with CheckUp ID: {checkup_id} and doc_id: {doc_id}")

        # Load and display data related to the checkup ID
        self.load_data()

        # Display lab tests in two frames
        self.display_lab_tests()
        self.ui.ProceedButton.setText("Update")

    def load_data(self):
        """Load both check-up and patient details and populate the UI."""
        try:
            # Step 1: Fetch check-up details
            checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            if not checkup_details:
                raise ValueError("No check-up details found for the given ID.")
            # Extract check-up data
            pat_id = checkup_details['pat_id']
            chck_bp = checkup_details['chck_bp']
            chck_temp = checkup_details['chck_temp']
            chck_height = checkup_details['chck_height']
            chck_weight = checkup_details['chck_weight']
            chckup_type = checkup_details['chckup_type']
            # Step 2: Fetch patient details
            patient_details = Patient.get_patient_details(pat_id)
            if not patient_details:
                raise ValueError("No patient details found for the given ID.")
            # Extract patient data
            pat_lname = patient_details['pat_lname']
            pat_fname = patient_details['pat_fname']
            pat_mname = patient_details['pat_mname']
            pat_dob = patient_details['pat_dob']
            pat_gender = patient_details['pat_gender']
            # Calculate age based on date of birth
            age = self.calculate_age(pat_dob)
            # Convert pat_dob to a string for display
            Birthday = pat_dob.strftime("%Y-%m-%d")
            # Step 3: Populate the UI
            self.populate_patient_info(pat_id, pat_lname, pat_fname, pat_mname, Birthday, age, pat_gender)
            self.populate_checkup_info(self.checkup_id, chck_bp, chck_temp, chck_height, chck_weight, chckup_type)

            # Store patient_id as an instance variable
            self.patient_id = pat_id  # Add this line

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load data: {e}")

    def calculate_age(self, dob):
        """Calculate the age based on the date of birth."""
        today = datetime.today()
        if isinstance(dob, date):  # Use the explicitly imported 'date'
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            return age
        else:
            raise ValueError("DOB must be a datetime.date object or a valid date string.")

    def populate_patient_info(self, pat_id , pat_lname, pat_fname, pat_mname, pat_dob, age, gender):
        """Populate the patient information fields."""
        self.ui.PatID.setText(str(pat_id))
        self.ui.PatName.setText(f"{pat_lname.capitalize()}, {pat_fname.capitalize()} {pat_mname.capitalize()}")
        self.ui.Dob.setText(pat_dob)
        self.ui.Age.setText(str(age))
        self.ui.PatGender.setText(gender)

    def populate_checkup_info(self, chck_id, chck_bp, chck_temp, chck_height, chck_weight, chckup_type):
        """Populate the check-up information fields."""
        self.ui.BP.setText(chck_bp + ' bpm')
        self.ui.Temperature.setText(chck_temp + ' °C')
        # Ensure the widget name matches the one in the .ui file
        if hasattr(self.ui, 'Height'):
            self.ui.Height.setText(str(chck_height) + ' cm')  # Convert to string and append unit
        else:
            QMessageBox.warning(self, "Missing Widget", "The 'Height' widget is missing in the UI.")
        self.ui.Weight.setText(str(chck_weight) + ' kg')
        self.ui.Type.setText(chckup_type)

    def display_lab_tests(self):
        """Display lab tests in two frames with pre-checked boxes for existing records."""
        try:
            # Fetch all lab tests
            tests = Laboratory.get_all_test()
            # Fetch lab codes already associated with the current check-up
            existing_lab_codes = CheckUp.get_lab_codes_by_chckid(self.checkup_id)

            # Debug: Log existing_lab_codes
            print(f"Existing lab codes for chck_id {self.checkup_id}: {existing_lab_codes}")

            # Convert existing_lab_codes to set for faster lookup
            existing_lab_codes_set = set(existing_lab_codes)

            # Count the total number of lab tests
            total_tests = len(tests)
            half = (total_tests + 1) // 2
            first_group = tests[:half]
            second_group = tests[half:]

            # Clear existing layouts in the frames
            self.clear_layout(self.ui.FirstFrame.layout())
            self.clear_layout(self.ui.SecondFrame.layout())

            # Add checkboxes to the first frame
            self.add_checkboxes_to_frame(first_group, self.ui.FirstFrame, existing_lab_codes_set)

            # Add checkboxes to the second frame
            self.add_checkboxes_to_frame(second_group, self.ui.SecondFrame, existing_lab_codes_set)

            # Connect the ProceedButton to process_selected_tests
            if hasattr(self.ui, 'ProceedButton'):
                self.ui.ProceedButton.clicked.connect(self.process_selected_tests)
            else:
                QMessageBox.warning(self, "Missing Button", "ProceedButton is missing in the UI.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to display lab tests: {e}")

    def ViewRecords(self):
        from Controllers.DoctorRecords_Controller import DoctorRecords
        print("Opening Doctor Records...")
        try:
            # Close the current DoctorDiagnosis window
            self.close()
            # Instantiate and show the DoctorRecords window with the doc_id
            self.doctor_records = DoctorRecords(doc_id=self.doc_id)
            self.doctor_records.show()
        except Exception as e:
            print(f"Error loading Doctor Records: {e}")
            QMessageBox.critical(self, "Error", f"Failed to load Doctor Records: {e}")

    def clear_layout(self, layout):
        """Clear all widgets from a layout."""
        if layout is not None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()

    def add_checkboxes_to_frame(self, tests, frame, existing_lab_codes_set):
        """Add checkboxes for the given tests to the specified frame and pre-check existing ones."""
        layout = frame.layout()
        if layout is None:
            layout = QVBoxLayout(frame)
            frame.setLayout(layout)

        for test in tests:
            lab_code = test["lab_code"]
            lab_name = test["lab_test_name"]

            # Create a QCheckBox for each test
            checkbox = QCheckBox(f"{lab_name}")
            checkbox.setProperty("lab_code", lab_code)

            # Pre-check the checkbox if the lab_code exists in the existing_lab_codes_set
            if lab_code in existing_lab_codes_set:
                checkbox.setChecked(True)

            layout.addWidget(checkbox)

    def process_selected_tests(self):
        """Process selected checkboxes and generate Lab Request PDF if applicable."""
        selected_lab_names = []

        for frame in [self.ui.FirstFrame, self.ui.SecondFrame]:
            layout = frame.layout()
            if layout:
                for i in range(layout.count()):
                    widget = layout.itemAt(i).widget()
                    if isinstance(widget, QCheckBox) and widget.isChecked():
                        lab_code = widget.property("lab_code")
                        if lab_code:
                            result = Laboratory.get_test_by_labcode(lab_code)
                            if result:
                                lab_name = result[0]['lab_test_name']
                                selected_lab_names.append(lab_name)

        # If no lab test selected, open modal
        if not selected_lab_names:
            self.open_doctor_lab_result_modal()
            return

        # Save raw lab codes only (excluding custom entries) to DB
        raw_lab_codes = [widget.property("lab_code")
                         for frame in [self.ui.FirstFrame, self.ui.SecondFrame]
                         for i in range(frame.layout().count())
                         if isinstance((widget := frame.layout().itemAt(i).widget()), QCheckBox)
                         and widget.isChecked()
                         and widget.property("lab_code")]

        success = CheckUp.update_lab_codes(self.checkup_id,
                                           raw_lab_codes )
        if not success:
            QMessageBox.critical(self, "Error", "Failed to update lab codes.")
            return

        try:
            # Get full patient and doctor info
            patient_info = Patient.get_patient_by_id(self.patient_id)
            doctor_info = Doctor.get_doctor(self.doc_id)
            if not patient_info or not doctor_info:
                QMessageBox.critical(self, "Error", "Failed to fetch patient or doctor information.")
                return

            name = self.ui.PatName.text()
            age = str(patient_info["age"])
            gender = patient_info["gender"]
            address = patient_info["address"]
            today = datetime.today().strftime("%Y-%m-%d")
            doctor_name = f"{doctor_info['first_name']} {doctor_info['middle_name']} {doctor_info['last_name']}" if doctor_info else "N/A"

            # Output paths
            output_dir = r"C:\Users\Roy Adrian Rondina\OneDrive - ctu.edu.ph\Desktop\Share"
            os.makedirs(output_dir, exist_ok=True)
            word_output = os.path.join(output_dir, f"temp_{self.checkup_id}_{name}.docx")
            pdf_output = os.path.join(output_dir, f"{self.checkup_id}_{name}_LabRequest.pdf")

            # Load the template
            template_path = r"C:\Users\Roy Adrian Rondina\PycharmProjects\IM-System\Images\LabRequest.docx"
            doc = Document(template_path)

            # Replace patient and doctor placeholders
            placeholders = {
                "{{name}}": name,
                "{{age}}": age,
                "{{gender}}": gender,
                "{{address}}": address,
                "{{date}}": today,
                "{{doctor_name}}": doctor_name
            }
            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    if key in p.text:
                        p.text = p.text.replace(key, val)

            # Fill in lab requests using test names
            for i in range(1, 11):
                tag = f"{{{{lab_request{i}}}}}"
                lab_name = f"• {selected_lab_names[i - 1]}" if i <= len(selected_lab_names) else ""
                for p in doc.paragraphs:
                    if tag in p.text:
                        p.text = p.text.replace(tag, lab_name)

            # Save and convert
            doc.save(word_output)
            convert(word_output, pdf_output)

            QMessageBox.information(self, "Success", f"PDF Lab Request created:\n{pdf_output}")
            self.close()

        except Exception as e:
            QMessageBox.critical(self, "PDF Error", f"Failed to generate Lab Request PDF: {e}")

