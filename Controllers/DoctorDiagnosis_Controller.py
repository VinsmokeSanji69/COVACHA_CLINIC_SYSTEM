from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMainWindow, QVBoxLayout, QCheckBox, QMessageBox, QApplication, QDialog, QLabel
from PyQt5.QtCore import Qt

from Controllers.ClientSocketController import DataRequest
from Controllers.DoctorLabResult_Controller import DoctorLabResult
from Views.Doctor_CheckUpList import Ui_Doctor_CheckUpList
from Views.Doctor_Diagnosis import Ui_Doctor_Diagnosis as DoctorDiagnosisUI
from Controllers.DoctorRecords_Controller import DoctorRecords
from datetime import datetime, date
from docx import Document
import os
from docx2pdf import convert
from Views.Doctor_Records import Ui_Doctor_Records

class DoctorDiagnosis(QMainWindow):
    def __init__(self, checkup_id, doc_id, parent=None):
        super().__init__(parent)
        self.ui = DoctorDiagnosisUI()
        self.ui.setupUi(self)

        # Store the checkup ID, doc_id, and reference to the DoctorDashboard
        self.checkup_id = checkup_id
        self.doc_id = doc_id

        # Load and display data related to the checkup ID
        self.load_data()

        # Display lab tests in two frames
        self.display_lab_tests()

    def load_data(self):
        """Load both check-up and patient details and populate the UI."""
        try:
            # Step 1: Fetch check-up details
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            if not checkup_details:
                raise ValueError("No check-up details found for the given ID.")

            # Extract check-up data
            pat_id = checkup_details['pat_id']
            chck_bp = checkup_details['chck_bp']
            chck_temp = checkup_details['chck_temp']
            chck_height = checkup_details['chck_height']
            chck_weight = checkup_details['chck_weight']
            chckup_type = checkup_details['chckup_type']

            # Step 2: Fetch patient details
            #patient_details = Patient.get_patient_details(pat_id)
            patient_details = DataRequest.send_command("GET_PATIENT_DETAILS", pat_id)

            if not patient_details:
                raise ValueError("No patient details found for the given ID.")

            # Extract patient data
            pat_lname = patient_details['pat_lname']
            pat_fname = patient_details['pat_fname']
            pat_mname = patient_details['pat_mname']
            pat_dob = patient_details['pat_dob']
            pat_gender = patient_details['pat_gender']

            # Calculate age based on date of birth
            age = self.calculate_age(pat_dob)

            # Convert pat_dob to a string for display
            Birthday = pat_dob.strftime("%Y-%m-%d")

            # Step 3: Populate the UI
            self.populate_patient_info(pat_id, pat_lname, pat_fname, pat_mname, Birthday, age, pat_gender)
            self.populate_checkup_info(self.checkup_id, chck_bp, chck_temp, chck_height, chck_weight, chckup_type)
            # Store patient_id as an instance variable
            self.patient_id = pat_id  # Add this line to store patient_id
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load data: {e}")

    def process_selected_tests(self):
        """Process selected checkboxes and generate Lab Request PDF if applicable."""
        selected_lab_names = []
        for frame in [self.ui.FirstFrame, self.ui.SecondFrame]:
            layout = frame.layout()
            if layout:
                for i in range(layout.count()):
                    widget = layout.itemAt(i).widget()
                    if isinstance(widget, QCheckBox) and widget.isChecked():
                        lab_code = widget.property("lab_code")
                        if lab_code:
                            #result = Laboratory.get_test_by_labcode(lab_code)
                            result = DataRequest.send_command("GET_TEST_BY_LAB_CODE", lab_code)

                            if result:
                                lab_name = result[0]['lab_test_name']
                                selected_lab_names.append(lab_name)

        if not selected_lab_names:
            self.ViewRecords()
            return

        raw_lab_codes = [widget.property("lab_code")
                         for frame in [self.ui.FirstFrame, self.ui.SecondFrame]
                         for i in range(frame.layout().count())
                         if isinstance((widget := frame.layout().itemAt(i).widget()), QCheckBox)
                         and widget.isChecked()
                         and widget.property("lab_code")]

        #success = CheckUp.update_lab_codes(self.checkup_id, raw_lab_codes)
        success = DataRequest.send_command("UPDATE_LAB_CODES",[self.checkup_id, raw_lab_codes])

        if not success:
            QMessageBox.critical(self, "Error", "Failed to update lab codes.")
            return

        try:
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            check_date_obj = checkup_details['chck_date']  # datetime.date
            folder_name = check_date_obj.strftime("%B %d, %Y")

            #patient_info = Patient.get_patient_by_id(self.patient_id)
            patient_info = DataRequest.send_command("GET_PATIENT_BY_ID", self.patient_id)

            #doctor_info = Doctor.get_doctor(self.doc_id)
            doctor_info = DataRequest.send_command("GET_DOCTOR_BY_ID", self.doc_id)

            if not patient_info or not doctor_info:
                QMessageBox.critical(self, "Error", "Failed to fetch patient or doctor information.")
                return

            name = self.ui.PatName.text()
            age = str(patient_info.get("age", ""))
            gender = patient_info.get("gender", "")
            address = patient_info.get("address", "")
            doctor_name = f"{doctor_info['first_name'].capitalize()} {doctor_info['middle_name'].capitalize()} {doctor_info['last_name'].capitalize()}"

            # Create folder
            base_dir = r"C:\Users\Roy Adrian Rondina\OneDrive - ctu.edu.ph\Desktop\Share"
            dated_folder_path = os.path.join(base_dir, folder_name)
            os.makedirs(dated_folder_path, exist_ok=True)

            pdf_output = os.path.join(dated_folder_path, f"{self.checkup_id}_{name}_LabRequest.pdf")
            word_output = os.path.join(dated_folder_path, f"temp_{self.checkup_id}_{name}.docx")

            template_path = r"Images\LabRequest.docx"
            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Template Error", "Lab Request template not found.")
                return

            # Show modal loading dialog that cannot be closed
            loading_dialog = QDialog(self)
            loading_dialog.setWindowTitle("Please Wait")
            loading_dialog.setWindowFlags(Qt.Window | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
            loading_dialog.setModal(True)
            layout = QVBoxLayout()
            layout.addWidget(QLabel("Generating Lab Request PDF, please wait..."))
            loading_dialog.setLayout(layout)
            loading_dialog.setFixedSize(300, 100)
            loading_dialog.show()
            QApplication.processEvents()

            # Generate Word document
            doc = Document(template_path)
            placeholders = {
                "{{name}}": name,
                "{{age}}": age,
                "{{gender}}": gender,
                "{{address}}": address,
                "{{date}}": check_date_obj.strftime("%B %d, %Y"),
                "{{doctor_name}}": doctor_name
            }

            for p in doc.paragraphs:
                for key, val in placeholders.items():
                    if key in p.text:
                        p.text = p.text.replace(key, val)

            for i in range(1, 11):
                tag = f"{{{{lab_request{i}}}}}"
                lab_name = f"• {selected_lab_names[i - 1]}" if i <= len(selected_lab_names) else ""
                for p in doc.paragraphs:
                    if tag in p.text:
                        p.text = p.text.replace(tag, lab_name)

            doc.save(word_output)

            try:
                convert(word_output, pdf_output)
                os.remove(word_output)
                loading_dialog.accept()
                QMessageBox.information(self, "Success", f"PDF Lab Request created:\n{pdf_output}")
            except Exception as conv_err:
                loading_dialog.accept()
                QMessageBox.warning(self, "Conversion Failed",
                                    f"Word document was saved but PDF conversion failed.\n\n"
                                    f"Reason: {conv_err}\n\n"
                                    f"Please ensure Microsoft Word is installed and the file is not open.")
                return

        except Exception as err:
            QMessageBox.critical(self, "PDF Error", f"Unexpected error while generating Lab Request:\n{err}")
            return

        self.ViewRecords()

    def calculate_age(self, dob):
        """Calculate the age based on the date of birth."""
        today = datetime.today()
        if isinstance(dob, date):  # Use the explicitly imported 'date'
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            return age
        else:
            raise ValueError("DOB must be a datetime.date object or a valid date string.")

    def populate_patient_info(self, pat_id , pat_lname, pat_fname, pat_mname, pat_dob, age, gender):
        """Populate the patient information fields."""
        self.ui.PatID.setText(str(pat_id))
        self.ui.PatName.setText(f"{pat_lname.capitalize()}, {pat_fname.capitalize()} {pat_mname.capitalize()}")
        self.ui.Dob.setText(pat_dob)
        self.ui.Age.setText(str(age))
        self.ui.PatGender.setText(gender)

    def populate_checkup_info(self, chck_id, chck_bp, chck_temp, chck_height, chck_weight, chckup_type):
        """Populate the check-up information fields."""
        self.ui.CheckID.setText(chck_id)
        self.ui.BP.setText(chck_bp + ' bpm')
        self.ui.Temperature.setText(chck_temp + ' °C')
        # Ensure the widget name matches the one in the .ui file
        if hasattr(self.ui, 'Height'):
            self.ui.Height.setText(str(chck_height) + ' cm')  # Convert to string and append unit
        else:
            QMessageBox.warning(self, "Missing Widget", "The 'Height' widget is missing in the UI.")
        self.ui.Weight.setText(str(chck_weight) + ' kg')
        self.ui.Type.setText(chckup_type)

    def display_lab_tests(self):
        """Display lab tests in two frames."""
        try:
            # Fetch all lab tests
            #tests = Laboratory.get_all_test()
            tests = DataRequest.send_command("GET_ALL_TEST")

            # Count the total number of lab tests
            #total_tests = Laboratory.count_all_test()
            total_tests = DataRequest.send_command("COUNT_ALL_TEST")

            # Divide the tests into two groups
            half = (total_tests + 1) // 2
            first_group = tests[:half]
            second_group = tests[half:]

            # Clear existing layouts in the frames
            self.clear_layout(self.ui.FirstFrame.layout())
            self.clear_layout(self.ui.SecondFrame.layout())

            # Add checkboxes to the first frame
            self.add_checkboxes_to_frame(first_group, self.ui.FirstFrame)

            # Add checkboxes to the second frame
            self.add_checkboxes_to_frame(second_group, self.ui.SecondFrame)

            # Connect the ProceedButton to process_selected_tests
            if hasattr(self.ui, 'ProceedButton'):
                self.ui.ProceedButton.clicked.connect(self.process_selected_tests)
            else:
                QMessageBox.warning(self, "Missing Button", "ProceedButton is missing in the UI.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to display lab tests: {e}")

    def ViewRecords(self):
        try:
            self.close()

            # Find DoctorDashboard window
            app = QApplication.instance()
            for widget in app.topLevelWidgets():
                if hasattr(widget, "page_stack") and hasattr(widget, "checkup_page"):
                    dashboard = widget
                    break
            else:
                return

            # Switch to the Checkup Page
            dashboard.page_stack.setCurrentWidget(dashboard.checkup_page)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load Doctor Records: {e}")

    def clear_layout(self, layout):
        """Clear all widgets from a layout."""
        if layout is not None:
            while layout.count():
                child = layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()

    def add_checkboxes_to_frame(self, tests, frame):
        """Add checkboxes for the given tests to the specified frame."""
        layout = frame.layout()
        if layout is None:
            layout = QVBoxLayout(frame)
            frame.setLayout(layout)

        checkbox_style = """
                        QCheckBox {
                            font-size: 14pt;
                            padding: 6px;
                        }
                        QCheckBox::indicator {
                            width: 30px;
                            height: 30px;
                            margin-right: 8px;
                        }
                        QCheckBox::indicator:checked {
                            background-color: #007acc;
                            border: none;
                        }
                        QCheckBox::indicator:unchecked {
                            background-color: transparent;
                            border: 2px solid #007acc;
                        }
                    """
        for test in tests:
            lab_code = test["lab_code"]
            lab_name = test["lab_test_name"]

            # Create a QCheckBox for each test
            checkbox = QCheckBox(f"{lab_name}")
            checkbox.setProperty("lab_code", lab_code)
            checkbox.setStyleSheet(checkbox_style)
            checkbox.setMinimumSize(35, 35)
            layout.addWidget(checkbox)

    def open_doctor_lab_result_modal(self):
        try:
            # Instantiate and show the DoctorLabResult modal
            self.doctor_lab_result = DoctorLabResult(checkup_id=self.checkup_id)
            self.doctor_lab_result.show()
            self.close()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open DoctorLabResult modal: {e}")