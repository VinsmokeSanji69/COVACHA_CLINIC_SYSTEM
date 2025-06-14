import os
import subprocess
import sys
from linecache import checkcache

from PyQt5 import QtWidgets, QtCore
from PyQt5.QtWidgets import QMainWindow, QVBoxLayout, QMessageBox, QLabel, QDialog, QDialogButtonBox, QApplication

from Controllers.ClientSocketController import DataRequest
from Views.Doctor_LabResult import Ui_Doctor_LabResult as DoctorLabResultUI
from Controllers.DoctorAddPrescription_Controller import DoctorAddPrescription
from datetime import datetime, date
from docx import Document
from docx2pdf import convert
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')  # Remove border
        tcBorders.append(border)

    tcPr.append(tcBorders)


class ConfirmationDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Confirm Diagnose")
        self.setFixedSize(400, 150)

        # Main layout
        layout = QVBoxLayout()

        # Add message label
        self.message_label = QLabel("Are you sure you want to proceed ?")
        layout.addWidget(self.message_label)

        # Add button box
        self.button_box = QDialogButtonBox(QDialogButtonBox.Yes | QDialogButtonBox.No)
        self.button_box.accepted.connect(self.accept)
        self.button_box.rejected.connect(self.reject)

        # Apply stylesheet to the button box
        self.button_box.setStyleSheet("""
            QPushButton {
                background-color: #2E6E65;
                color: white;
                border-radius: 10px;
                padding: 5px 10px;
                margin-top: 5px
            }
            QPushButton:hover {
                background-color: #235C5A;
            }
        """)
        layout.addWidget(self.button_box)

        # Set layout
        self.setLayout(layout)

class DoctorLabResult(QMainWindow):
    def __init__(self, checkup_id=None, parent=None, refresh_callback=None, view = False):
        super().__init__(parent)
        self.ui = DoctorLabResultUI()
        self.ui.setupUi(self)
        self.checkup_id = checkup_id
        self.refresh_callback = refresh_callback

        # Load and display data related to the checkup ID
        self.load_data()
        self.apply_table_styles()
        self.refresh_all_tables()
        self.view = view

        # Connect buttons
        self.ui.ViewLabResult.clicked.connect(self.view_file)


        if self.view is True:
            self.setup_view()
            self.initialize_diagnosis()
        else:
            self.ui.AddPrescription.clicked.connect(self.open_add_prescription_form)
            self.ui.DiagnoseButton.clicked.connect(self.confirm_and_add_diagnosis)
            self.ui.Cancel.clicked.connect(self.return_to_dashboard)
            self.ui.EditPrescription.clicked.connect(self.open_edit_form)
            self.ui.ViewLabResult_4.clicked.connect(self.delete_prescription)

    def setup_view(self):
        self.ui.AddPrescription.setVisible(False)
        self.ui.EditPrescription.setVisible(False)
        self.ui.ViewLabResult_4.setVisible(False)
        self.ui.DiagnoseButton.setVisible(False)
        self.ui.Cancel.setVisible(False)
        self.ui.DiagnoseText.setReadOnly(True)
        self.ui.PrescriptionLabel.setReadOnly(True)


    def initialize_diagnosis(self):
        #checkup = CheckUp.get_checkup_details(self.checkup_id)
        checkup = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

        diagnosis = checkup.get("chck_diagnoses")
        prescnotes = checkup.get("chck_notes")
        self.ui.DiagnoseText.setText(diagnosis)
        self.ui.PrescriptionLabel.setText(prescnotes)

    def return_to_dashboard(self):
        try:
            # Refresh the parent window if a refresh callback is provided
            if hasattr(self, 'refresh_callback') and callable(self.refresh_callback):
                self.refresh_callback()

            # Get the parent window (dashboard)
            parent_window = self.parent()

            # Close the modal
            self.close()

            # Bring the parent window back into focus
            if parent_window:
                parent_window.activateWindow()
                parent_window.raise_()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while closing the modal: {e}")

    def refresh_all_tables(self):
        try:
            self.load_labattach_table()
            self.load_prescription_table()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to refresh tables: {e}")

    def apply_table_styles(self):
        self.ui.LabTestTabe.setStyleSheet("""
                                  QTableWidget {
                       background-color: #F4F7ED;
                       gridline-color: transparent;
                       border-radius: 10px;
                   }
                   QTableWidget::item {
                       border: none;
                       font: 16pt "Lexend";
                   }
                   QTableWidget::item:selected {
                       background-color: rgba(46, 110, 101, 0.3);
                   }
                   QTableWidget QHeaderView::section {
                       background-color: #2E6E65;
                       color: white;
                       padding: 5px;
                       font: 18px "Lexend Medium";
                       border: 2px solid #2E6E65;
                   }

                   Scroll Area CSS
                   QScrollBar:vertical {
                        background: transparent;
                        width: 10px;
                       border-radius: 5px;
                   }
                   QScrollBar::handle:vertical {
                           background: #C0C0C0;
                           border-radius: 5px;
                   }
                   QScrollBar::handle:vertical:hover {
                           background: #A0A0A0;
                   }
                   QScrollBar::add-line:vertical,
                   QScrollBar::sub-line:vertical{
                           background: none;
                           border: none;
                   }
                      """)
        self.ui.LabTestTabe.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        # Ensure horizontal headers are visible
        self.ui.LabTestTabe.horizontalHeader().setVisible(True)

        # Align headers to the left
        self.ui.LabTestTabe.horizontalHeader().setDefaultAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)

        # Hide the vertical header (row index)
        self.ui.LabTestTabe.verticalHeader().setVisible(False)

        self.ui.LabTestTabe_2.setStyleSheet("""
                                          QTableWidget {
                               background-color: #F4F7ED;
                               gridline-color: transparent;
                               border-radius: 10px;
                           }
                           QTableWidget::item {
                               border: none;
                               font: 16pt "Lexend";
                           }
                           QTableWidget::item:selected {
                               background-color: rgba(46, 110, 101, 0.3);
                           }
                           QTableWidget QHeaderView::section {
                               background-color: #2E6E65;
                               color: white;
                               padding: 5px;
                               font: 18px "Lexend Medium";
                               border: 2px solid #2E6E65;
                           }

                           Scroll Area CSS
                           QScrollBar:vertical {
                                background: transparent;
                                width: 10px;
                               border-radius: 5px;
                           }
                           QScrollBar::handle:vertical {
                                   background: #C0C0C0;
                                   border-radius: 5px;
                           }
                           QScrollBar::handle:vertical:hover {
                                   background: #A0A0A0;
                           }
                           QScrollBar::add-line:vertical,
                           QScrollBar::sub-line:vertical{
                                   background: none;
                                   border: none;
                           }
                              """)
        self.ui.LabTestTabe_2.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        # Ensure horizontal headers are visible
        self.ui.LabTestTabe_2.horizontalHeader().setVisible(True)

        # Align headers to the left
        self.ui.LabTestTabe_2.horizontalHeader().setDefaultAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)

        # Hide the vertical header (row index)
        self.ui.LabTestTabe_2.verticalHeader().setVisible(False)

    def load_labattach_table(self):
        try:
            # Clear the table before populating it
            self.ui.LabTestTabe.setRowCount(0)

            # Fetch lab codes and attachments for the given check-up ID
            #lab_tests = CheckUp.get_test_names_by_chckid(self.checkup_id)
            lab_tests = DataRequest.send_command("GET_TEST_BY_CHECK_ID", self.checkup_id)

            if not lab_tests:
                # Add a single row with "No Lab Test Request"
                row_position = self.ui.LabTestTabe.rowCount()
                self.ui.LabTestTabe.insertRow(row_position)
                self.ui.LabTestTabe.setItem(row_position, 0, QtWidgets.QTableWidgetItem("No Lab Test Request"))
                self.ui.LabTestTabe.setItem(row_position, 1, QtWidgets.QTableWidgetItem(""))  # Empty attachment status
                return

            # Populate the table with lab test details
            for lab_test in lab_tests:
                # Handle both dictionaries and tuples
                if isinstance(lab_test, dict):
                    lab_code = lab_test.get('lab_code')
                    lab_attachment = lab_test.get('lab_attachment')
                else:
                    continue

                # Validate lab_code
                if not lab_code:
                    continue

                # Fetch the lab test name using the lab code
                #lab_test_details = Laboratory.get_test_by_labcode(lab_code)
                lab_test_details = DataRequest.send_command("GET_TEST_BY_LAB_CODE", lab_code)

                if not lab_test_details:
                    continue

                # Handle the case where get_test_by_labcode returns a tuple
                if isinstance(lab_test_details, list):
                    # Assume the first element of the tuple contains 'lab_test_name'
                    if len(lab_test_details) > 0 and isinstance(lab_test_details[0], dict):
                        lab_test_details = lab_test_details[0]
                    else:
                        continue

                # Validate lab_test_details
                if not isinstance(lab_test_details, dict):
                    continue

                lab_test_name = lab_test_details.get('lab_test_name', 'Unknown').capitalize()

                # Determine attachment status
                if lab_attachment:
                    # Convert memoryview to string and extract the file name
                    if isinstance(lab_attachment, memoryview):
                        lab_attachment = lab_attachment.tobytes().decode('utf-8')
                    file_name = os.path.basename(lab_attachment)
                    attachment_status = file_name
                else:
                    attachment_status = "No Attach File"

                # Add a new row to the table
                row_position = self.ui.LabTestTabe.rowCount()
                self.ui.LabTestTabe.insertRow(row_position)

                # Insert data into the table
                self.ui.LabTestTabe.setItem(row_position, 0, QtWidgets.QTableWidgetItem(lab_test_name))
                self.ui.LabTestTabe.setItem(row_position, 1, QtWidgets.QTableWidgetItem(attachment_status))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load lab attach table: {e}")

    def load_prescription_table(self):
        """Load and display prescription data in the LabTestTabe_2 table."""
        try:
            # Clear existing rows in the table
            self.ui.LabTestTabe_2.setRowCount(0)

            # Fetch prescriptions for the given check-up ID
            #prescriptions = Prescription.display_prescription(self.checkup_id)
            prescriptions = DataRequest.send_command("GET_PRESCRIPTION_BY_CHECKUP",self.checkup_id)
            if not prescriptions:
                # Add a single row with "No Prescriptions"
                row_position = self.ui.LabTestTabe_2.rowCount()
                self.ui.LabTestTabe_2.insertRow(row_position)
                self.ui.LabTestTabe_2.setItem(row_position, 0, QtWidgets.QTableWidgetItem("No Prescriptions"))
                return

            # Populate the table with prescription details
            for prescription in prescriptions:
                # Extract prescription data
                med_name = prescription.get("pres_medicine", "")
                dosage = prescription.get("pres_dosage", "")
                intake = prescription.get("pres_intake", "")
                tablets = prescription.get("pres_tablets", "")

                # Add a new row to the table
                row_position = self.ui.LabTestTabe_2.rowCount()
                self.ui.LabTestTabe_2.insertRow(row_position)

                # Insert data into the table
                self.ui.LabTestTabe_2.setItem(row_position, 0, QtWidgets.QTableWidgetItem(med_name))
                self.ui.LabTestTabe_2.setItem(row_position, 1, QtWidgets.QTableWidgetItem(dosage))
                self.ui.LabTestTabe_2.setItem(row_position, 2, QtWidgets.QTableWidgetItem(intake))
                self.ui.LabTestTabe_2.setItem(row_position, 3, QtWidgets.QTableWidgetItem(tablets))

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load prescription table: {e}")

    def load_data(self):
        """Load both check-up and patient details and populate the UI."""
        try:
            # Step 1: Fetch check-up details
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            if not checkup_details:
                raise ValueError("No check-up details found for the given ID.")

            # Extract check-up data
            pat_id = checkup_details['pat_id']
            chck_bp = checkup_details['chck_bp']
            chck_temp = checkup_details['chck_temp']
            chck_height = checkup_details['chck_height']
            chck_weight = checkup_details['chck_weight']

            # Step 2: Fetch patient details
            #patient_details = Patient.get_patient_details(pat_id)
            patient_details = DataRequest.send_command("GET_PATIENT_DETAILS", pat_id)

            if not patient_details:
                raise ValueError("No patient details found for the given ID.")

            # Extract patient data
            pat_lname = patient_details['pat_lname']
            pat_fname = patient_details['pat_fname']
            pat_mname = patient_details['pat_mname']
            pat_dob = patient_details['pat_dob']
            pat_gender = patient_details['pat_gender']
            pat_contact = patient_details['pat_contact']

            # Calculate age based on date of birth
            age = self.calculate_age(pat_dob)

            # Convert pat_dob to a string for display
            Birthday = pat_dob.strftime("%Y-%m-%d")

            # Step 3: Populate the UI
            self.populate_patient_info(pat_id, pat_lname, pat_fname, pat_mname, Birthday, age, pat_gender)
            self.populate_checkup_info(chck_bp, chck_temp, chck_height, chck_weight)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load data: {e}")

    def calculate_age(self, dob):
        """Calculate the age based on the date of birth."""
        today = datetime.today()
        if isinstance(dob, date):  # Use the explicitly imported 'date'
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            return age
        else:
            raise ValueError("DOB must be a datetime.date object or a valid date string.")

    def populate_patient_info(self, pat_id, pat_lname, pat_fname, pat_mname, pat_dob, age, gender):
        """Populate the patient information fields."""
        self.ui.PatID.setText(str(pat_id))
        self.ui.PatName.setText(f"{pat_lname.capitalize()}, {pat_fname.capitalize()} {pat_mname.capitalize()}")
        self.ui.PatDoB.setText(pat_dob)
        self.ui.PatAge.setText(str(age))
        self.ui.PatGender.setText(gender)

    def populate_checkup_info(self, chck_bp, chck_temp, chck_height, chck_weight):
        """Populate the check-up information fields."""
        self.ui.BloodPressure.setText(str(chck_bp + " bpm"))
        self.ui.Temperature.setText(str(chck_temp + " °C"))
        self.ui.Heights.setText(str(chck_height + " cm"))
        self.ui.Weight.setText(str(chck_weight + " kg"))
        self.ui.CheckUpID.setText(self.checkup_id)

    def view_file(self):
        """Handle viewing the attached file for the selected lab test."""
        # Get the currently selected row in the LabTable
        selected_row = self.ui.LabTestTabe.currentRow()
        if selected_row == -1:  # No row selected
            QMessageBox.warning(self, "Selection Error", "Please select a lab test from the table.")
            return

        # Retrieve the lab_test_name from the selected row
        lab_test_name = self.ui.LabTestTabe.item(selected_row, 0).text()

        # Normalize the lab_test_name: strip whitespace and convert to lowercase
        lab_test_name = lab_test_name.strip().lower()

        # Retrieve the lab_code using the normalized lab_test_name
        #lab_code = Laboratory.get_lab_code_by_name(lab_test_name)
        lab_code = DataRequest.send_command("GET_LAB_CODE_BY_NAME", lab_test_name)

        if not lab_code:
            QMessageBox.critical(self, "Error", "Failed to retrieve lab code.")
            return
        # Fetch the file path from the CheckUp model
        #file_path = CheckUp.get_lab_attachment(self.checkup_id, lab_code)
        file_path = DataRequest.send_command("GET_LAB_ATTACHMENT", [self.checkup_id, lab_code])

        if not file_path:
            QMessageBox.warning(self, "No Attachment", "No file is attached to this lab test.")
            return

        # Check if the file exists
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' does not exist.")
            return

        # Open the file using the default application
        try:
            if os.name == 'nt':  # Windows
                os.startfile(file_path)
            else:  # Unix-like systems
                opener = "open" if sys.platform == "darwin" else "xdg-open"
                subprocess.call([opener, file_path])
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open file: {e}")

    def open_edit_form(self):
        """Open the DoctorAddPrescription form with pre-filled prescription data."""
        selected_row = self.ui.LabTestTabe_2.currentRow()
        if selected_row == -1:
            QMessageBox.warning(self, "Selection Error", "Please select a prescription to edit.")
            return

        # Get prescription data from the selected row
        med_name = self.ui.LabTestTabe_2.item(selected_row, 0).text()
        dosage = self.ui.LabTestTabe_2.item(selected_row, 1).text()
        intake = self.ui.LabTestTabe_2.item(selected_row, 2).text()
        tablets = self.ui.LabTestTabe_2.item(selected_row, 3).text()

        # Fetch prescription by details
        #pres_data = Prescription.get_prescription_by_details(self.checkup_id, med_name, dosage, intake, tablets)
        pres_data = DataRequest.send_command("GET_PRESCRIPTION_BY_DETAILS", [self.checkup_id, med_name, dosage, intake, tablets])

        if not pres_data:
            QMessageBox.critical(self, "Error", "Failed to find prescription in database.")
            return

        try:
            # Only now proceed to create the edit window
            self.edit_prescription_window = DoctorAddPrescription(
                chck_id=self.checkup_id,
                parent=self,
                refresh_callback=self.refresh_all_tables,
                prescription_data=pres_data
            )
            self.edit_prescription_window.show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open edit form: {e}")

    def delete_prescription(self):
        selected_row = self.ui.LabTestTabe_2.currentRow()
        if selected_row == -1:
            QMessageBox.warning(self, "Selection Error", "Please select a prescription to delete.")
            return

        # Show custom confirmation dialog
        dlg = ConfirmationDialog(self)
        if dlg.exec_() == QDialog.Rejected:
            return

        # Proceed with deletion
        med_name = self.ui.LabTestTabe_2.item(selected_row, 0).text()
        dosage = self.ui.LabTestTabe_2.item(selected_row, 1).text()
        intake = self.ui.LabTestTabe_2.item(selected_row, 2).text()

        #pres_data = Prescription.get_prescription_by_details(self.checkup_id, med_name, dosage, intake)
        pres_data = DataRequest.send_command("GET_PRESCRIPTION_BY_DETAILS", [self.checkup_id, med_name, dosage, intake])

        if not pres_data:
            QMessageBox.critical(self, "Error", "Failed to find prescription in database.")
            return

        if DataRequest.send_command("DELETE_PRESCRIPTION",pres_data['pres_id']):
            QMessageBox.information(self, "Success", "Prescription deleted successfully.")
            self.refresh_all_tables()
        else:
            QMessageBox.critical(self, "Error", "Failed to delete prescription.")

    def open_add_prescription_form(self):
        self.add_prescription_window = DoctorAddPrescription(
            chck_id=self.checkup_id,
            parent=self,
            refresh_callback=self.refresh_all_tables
        )
        self.add_prescription_window.show()

    def confirm_and_add_diagnosis(self):
        try:
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            # Get diagnosis text and notes from the UI
            chck_diagnoses = self.ui.DiagnoseText.toPlainText().strip()
            chck_notes = self.ui.PrescriptionLabel.toPlainText().strip()
            # Validate diagnosis text
            if not chck_diagnoses:
                QMessageBox.warning(self, "Validation Error", "Diagnosis text is required.")
                return

            # Show confirmation dialog
            confirmation_dialog = ConfirmationDialog(self)
            if confirmation_dialog.exec_() == QDialog.Rejected:
                return

            # Update check-up status to "Completed"
            #success = CheckUp.change_status_completed(self.checkup_id)
            success = DataRequest.send_command("CHANGE_STATUS_COMPLETED",self.checkup_id)
            if not success:
                raise ValueError("Failed to change check-up status to Completed.")
            chck_id = self.checkup_id,
            chck_diagnoses=chck_diagnoses,
            chck_notes=chck_notes
            # Save diagnosis notes
            # success = CheckUp.add_diagnosis_notes(
            #     chck_id=self.checkup_id,
            #     chck_diagnoses=chck_diagnoses,
            #     chck_notes=chck_notes
            # )
            success = DataRequest.send_command("ADD_DIAGNOSIS_NOTES",[chck_id, chck_diagnoses,chck_notes])
            if not success:
                raise ValueError("Failed to save diagnosis notes.")

            # Notify user of success
            QMessageBox.information(self, "Success", "Diagnosis saved successfully!")

            pat_id = checkup_details['pat_id']

            # Show "Generating PDF..." message
            pdf_progress = QMessageBox(self)
            pdf_progress.setIcon(QMessageBox.Information)
            pdf_progress.setWindowTitle("Please Wait")
            pdf_progress.setText("Generating PDF files...")
            pdf_progress.setStandardButtons(QMessageBox.NoButton)
            pdf_progress.show()

            # Ensure the dialog shows up immediately
            QApplication.processEvents()

            # Generate PDFs (this may block the UI slightly, but it's expected)
            self.make_into_pdf(pat_id)
            self.make_prescription_pdf(pat_id)

            # Hide the "Generating PDF" dialog
            pdf_progress.hide()  # Use hide() instead of close()

            # Small delay to ensure the dialog is fully hidden (optional)
            time.sleep(0.1)  # Adjust as needed

            # Show success message
            pdf_success = QMessageBox(self)
            pdf_success.setIcon(QMessageBox.Information)
            pdf_success.setWindowTitle("PDF Generated")
            pdf_success.setText("PDF files have been generated successfully.")
            pdf_success.setStandardButtons(QMessageBox.Ok)
            pdf_success.exec_()

            # Refresh the tables in the parent window
            if self.refresh_callback:
                self.refresh_callback()

            # Close current window
            self.close()

            # Close parent window if exists
            parent_window = self.parent()
            if parent_window:
                parent_window.close()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {e}")

    def make_into_pdf(self, pat_id):
        try:
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            #patient_details = Patient.get_patient_by_id(pat_id)
            patient_details = DataRequest.send_command("GET_PATIENT_BY_ID", pat_id)

            doc_id = checkup_details['doc_id']
            #doctor_details = Doctor.get_doctor(doc_id)
            doctor_details = DataRequest.send_command("GET_DOCTOR_BY_ID", doc_id)


            doctor_name = f"{doctor_details['first_name']} {doctor_details['middle_name']} {doctor_details['last_name']}" if doctor_details else "N/A"

            # Prepare data dictionary matching your template placeholders
            data = {
                'name': self.ui.PatName.text() or '___',
                'age': self.ui.PatAge.text() or '___',
                'gender': self.ui.PatGender.text() or '___',
                'address': patient_details.get('address', '___'),
                'contact': patient_details.get('contact', '___'),
                'dob': patient_details.get('dob', '___'),
                'date': checkup_details['chck_date'] or '___',
                'doctor_name': doctor_name,
                'bloodpressure': self.ui.BloodPressure.text() or '___',
                'temperature': self.ui.Temperature.text() or '___',
                'weight': self.ui.Weight.text() or '___',
                'height': self.ui.Heights.text() or '___',
                'diagnosis': self.ui.DiagnoseText.toPlainText() or '___',
            }

            # Ensure check_date_raw is a string
            check_date_raw = checkup_details['chck_date']
            if isinstance(check_date_raw, date):  # Use 'date' instead of 'datetime.date'
                check_date_raw = check_date_raw.strftime("%Y-%m-%d")  # Convert to string

            # Format the checkup date to "June 09, 2025"
            check_date_obj = datetime.strptime(check_date_raw, "%Y-%m-%d")
            folder_name = check_date_obj.strftime("%B %d, %Y")

            # Create a folder for that date
            base_dir = r"C:\Users\Roy Adrian Rondina\OneDrive - ctu.edu.ph\Desktop\Share"
            dated_folder_path = os.path.join(base_dir, folder_name)
            os.makedirs(dated_folder_path, exist_ok=True)

            # Set file paths inside the dated folder
            word_output = os.path.join(dated_folder_path, f"temp_{self.checkup_id}_{data['name']}.docx")
            pdf_output = os.path.join(dated_folder_path, f"{self.checkup_id}_{data['name']}_Diagnosis.pdf")

            # Load and fill template
            doc = Document("Images/PatientRecord.docx")

            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))

            doc.save(word_output)
            convert(word_output, pdf_output)
            os.remove(word_output)

            return pdf_output

        except Exception as e:
            error_msg = f"Error generating PDF: {str(e)}"
            QMessageBox.critical(self, "PDF Generation Error", error_msg)
            return None

    def make_prescription_pdf(self, pat_id):
        try:
            #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
            checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

            #patient_details = Patient.get_patient_by_id(pat_id)
            patient_details = DataRequest.send_command("GET_PATIENT_BY_ID", pat_id)

            #doctor_details = Doctor.get_doctor(checkup_details['doc_id'])
            doctor_details = DataRequest.send_command("GET_DOCTOR_BY_ID", checkup_details['doc_id'])

            #prescriptions = Prescription.display_prescription(self.checkup_id)
            prescriptions = DataRequest.send_command("GET_PRESCRIPTION_BY_CHECKUP",self.checkup_id)

            doctor_name = f"{doctor_details['first_name']} {doctor_details['middle_name']} {doctor_details['last_name']}" if doctor_details else "N/A"

            data = {
                'name': self.ui.PatName.text() or '___',
                'age': self.ui.PatAge.text() or '___',
                'gender': self.ui.PatGender.text() or '___',
                'address': patient_details.get('address', '___'),
                'date': checkup_details['chck_date'] or '___',
                'prescription_notes': self.ui.PrescriptionLabel.toPlainText().strip() or 'None',
                'doctor_name': doctor_name
            }

            check_date_raw = checkup_details['chck_date']
            if isinstance(check_date_raw, date):
                check_date_obj = check_date_raw
            else:
                check_date_obj = datetime.strptime(check_date_raw, "%Y-%m-%d")
            folder_name = check_date_obj.strftime("%B %d, %Y")

            base_dir = r"C:\Users\Roy Adrian Rondina\OneDrive - ctu.edu.ph\Desktop\Share"
            dated_folder_path = os.path.join(base_dir, folder_name)
            os.makedirs(dated_folder_path, exist_ok=True)

            filename_base = f"{self.checkup_id}_{data['name']}_Prescription"
            word_output = os.path.join(dated_folder_path, f"temp_{filename_base}.docx")
            pdf_output = os.path.join(dated_folder_path, f"{filename_base}.pdf")

            template_path = r"Images\Prescription.docx"
            doc = Document(template_path)

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            # Fill prescription table dynamically
            for table in doc.tables:
                if table.cell(0, 0).text.strip().lower().startswith("medicine"):
                    while len(table.rows) > 1:
                        table._tbl.remove(table.rows[1]._tr)

                    for pres in prescriptions:
                        row_cells = table.add_row().cells
                        row_cells[0].text = pres.get("pres_medicine", "")
                        row_cells[1].text = pres.get("pres_dosage", "")
                        row_cells[2].text = pres.get("pres_tablets", "")
                        row_cells[3].text = pres.get("pres_intake", "")

                        # Remove borders from each cell in the new row
                        for cell in row_cells:
                            remove_cell_borders(cell)
                    break

            doc.save(word_output)
            convert(word_output, pdf_output)
            os.remove(word_output)
            return pdf_output

        except Exception as e:
            error_msg = f"Error generating prescription PDF: {str(e)}"
            QMessageBox.critical(self, "PDF Generation Error", error_msg)
            return None

    def open_or_focus_doctor_records(self):
        from Controllers.DoctorRecords_Controller import DoctorRecords
        app = QApplication.instance()

        # Check if any DoctorRecords window is already active
        for widget in app.topLevelWidgets():
            if isinstance(widget, DoctorRecords):
                doctor_records_window = widget
                doctor_records_window.activateWindow()
                doctor_records_window.show()
                return

        # If no existing window is found, create a new one
        #checkup_details = CheckUp.get_checkup_details(self.checkup_id)
        checkup_details = DataRequest.send_command("GET_CHECKUP_DETAILS", self.checkup_id)

        if not checkup_details or 'doc_id' not in checkup_details:
            return

        doc_id = checkup_details['doc_id']

        # Create a new DoctorRecords window and store it as an instance variable
        self.doctor_records_window = DoctorRecords(doc_id=doc_id)
        self.doctor_records_window.show()
